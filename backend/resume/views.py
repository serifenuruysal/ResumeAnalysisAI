import logging
import openai
from django.conf import settings
from django.conf import settings
from rest_framework.decorators import api_view, parser_classes  # Import parser_classes
from rest_framework.parsers import MultiPartParser, FormParser  # Import necessary parsers
from rest_framework.response import Response
from rest_framework import status
from PyPDF2 import PdfReader
import docx

# Configure logging
logger = logging.getLogger(__name__)
openai.api_key = settings.OPENAI_API_KEY

@api_view(['POST'])
def upload_and_analyze(request):
    file = request.FILES.get('file')

    if not file:
        logger.warning("No file uploaded.")
        return Response({'error': 'No file uploaded.'}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        # Extract text from PDF, DOCX, or TXT files
        resume_text = ''
        if file.name.endswith('.pdf'):
            reader = PdfReader(file)
            resume_text = ''.join([page.extract_text() for page in reader.pages])
            logger.info(f"Extracted text from PDF: {resume_text[:100]}...")  # Log first 100 chars
        elif file.name.endswith('.docx'):
            doc = docx.Document(file)
            resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"Extracted text from DOCX: {resume_text[:100]}...")  # Log first 100 chars
        elif file.name.endswith('.txt'):
            resume_text = file.read().decode('utf-8')
            logger.info(f"Extracted text from TXT: {resume_text[:100]}...")  # Log first 100 chars
        else:
            logger.warning("Unsupported file format.")
            return Response({'error': 'Unsupported file format.'}, status=status.HTTP_400_BAD_REQUEST)

        logger.info("Resume uploaded and text extracted successfully.")

        # Analyze the extracted text using OpenAI
        prompt = f"Analyze the following resume and provide feedback on its strengths and areas of improvement:\n\n{resume_text}"
        logger.info("Sending request to OpenAI for analysis.")

        response = openai.ChatCompletion.create(
            model="gpt-3.5-turbo",  # or "gpt-4" depending on your needs
            messages=[{"role": "user", "content": prompt}],
            max_tokens=500
        )

        analysis = response.choices[0].message['content'].strip()
        logger.info({"Resume analysis completed successfully.",analysis})
        
        return Response({'analysis': analysis}, status=status.HTTP_200_OK)

    except Exception as e:
        logger.exception("An error occurred during resume processing.")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
    
@api_view(['POST'])
def upload_resume(request):
    file = request.FILES.get('file')

    if not file:
        logger.warning("No file uploaded.")
        return Response({'error': 'No file uploaded.'}, status=status.HTTP_400_BAD_REQUEST)
    
    try:
        # Process PDF, DOCX, or TXT files
        resume_text = ''
        if file.name.endswith('.pdf'):
            reader = PdfReader(file)
            resume_text = ''.join([page.extract_text() for page in reader.pages])
            logger.info(f"Extracted text from PDF: {resume_text[:100]}...")  # Log first 100 chars
        elif file.name.endswith('.docx'):
            doc = docx.Document(file)
            resume_text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])
            logger.info(f"Extracted text from DOCX: {resume_text[:100]}...")  # Log first 100 chars
        elif file.name.endswith('.txt'):
            resume_text = file.read().decode('utf-8')
            logger.info(f"Extracted text from TXT: {resume_text[:100]}...")  # Log first 100 chars
        else:
            logger.warning("Unsupported file format.")
            return Response({'error': 'Unsupported file format.'}, status=status.HTTP_400_BAD_REQUEST)

        logger.info("Resume uploaded and text extracted successfully.")
        return Response({'resume_text': resume_text}, status=status.HTTP_200_OK)

    except Exception as e:
        logger.exception("An error occurred while processing the resume upload.")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

@api_view(['POST'])
def analyze_resume(request):
    resume_text = request.data.get('resume_text')

    if not resume_text:
        logger.warning("No resume text provided.")
        return Response({'error': 'No resume text provided.'}, status=status.HTTP_400_BAD_REQUEST)

    try:
        # Use OpenAI to analyze the resume text
        prompt = f"Analyze the following resume and provide feedback on its strengths and areas of improvement:\n\n{resume_text}"
        logger.info("Sending request to OpenAI for analysis.")

        response = openai.Completion.create(
            engine="gpt-4",
            prompt=prompt,
            max_tokens=500
        )

        analysis = response.choices[0].text.strip()
        logger.info("Resume analysis completed successfully.")
        return Response({'analysis': analysis}, status=status.HTTP_200_OK)

    except Exception as e:
        logger.exception("An error occurred during resume analysis.")
        return Response({'error': str(e)}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)
